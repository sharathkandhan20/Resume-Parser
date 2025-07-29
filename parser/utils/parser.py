"""
Enhanced Resume Parser with Gemini API integration
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, Optional, List, Any
import io
import re
import time
import threading
from datetime import datetime, timedelta
from collections import defaultdict
from difflib import SequenceMatcher

# Document processing libraries
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import pandas as pd
import numpy as np

# API
import google.generativeai as genai
from dotenv import load_dotenv

# Configure logging
logger = logging.getLogger(__name__)


# ✅ 1. SMART GEMINI API KEY MANAGER
class GeminiKeyManager:
    """Manages multiple Gemini API keys with rate limiting and automatic rotation"""
    
    def __init__(self):
        self.keys = self._load_all_keys()
        self.key_usage = defaultdict(lambda: {
            'requests_per_minute': [],
            'requests_today': 0,
            'tokens_per_minute': [],
            'last_reset': datetime.now().date(),
            'is_exhausted': False
        })
        self.lock = threading.Lock()
        
        # Gemini free-tier limits
        self.RPM_LIMIT = 15
        self.RPD_LIMIT = 1500
        self.TPM_LIMIT = 1000000
        
        logger.info(f"Initialized GeminiKeyManager with {len(self.keys)} keys")
    
    def _load_all_keys(self) -> List[str]:
        """Load all available Gemini API keys from environment"""
        keys = []
        
        # Try numbered keys (GEMINI_API_KEY_1, GEMINI_API_KEY_2, etc.)
        for i in range(1, 20):  # Support up to 20 keys
            key = os.getenv(f'GEMINI_API_KEY_{i}')
            if key:
                keys.append(key)
                logger.debug(f"Loaded GEMINI_API_KEY_{i}")
        
        # Also try single key format
        single_key = os.getenv('GEMINI_API_KEY')
        if single_key and single_key not in keys:
            keys.append(single_key)
            logger.debug("Loaded GEMINI_API_KEY")
        
        return keys
    
    def estimate_tokens(self, text: str) -> int:
        """Estimate token count for text (rough approximation)"""
        # Rough estimate: 1 token ≈ 0.75 words
        word_count = len(text.split())
        return int(word_count * 1.3)
    
    def _clean_old_usage(self, key: str):
        """Remove usage records older than 1 minute and reset daily counter if needed"""
        now = datetime.now()
        usage = self.key_usage[key]
        
        # Clean requests older than 1 minute
        usage['requests_per_minute'] = [
            req_time for req_time in usage['requests_per_minute']
            if now - req_time < timedelta(minutes=1)
        ]
        
        # Clean token usage older than 1 minute
        usage['tokens_per_minute'] = [
            (token_time, tokens) for token_time, tokens in usage['tokens_per_minute']
            if now - token_time < timedelta(minutes=1)
        ]
        
        # Reset daily counter if it's a new day
        if usage['last_reset'] != now.date():
            usage['requests_today'] = 0
            usage['last_reset'] = now.date()
            usage['is_exhausted'] = False
            logger.info(f"Reset daily counter for key ending in ...{key[-4:]}")
    
    def get_available_key(self, estimated_tokens: int) -> Optional[str]:
        """Get an available API key that hasn't hit rate limits"""
        with self.lock:
            for key in self.keys:
                self._clean_old_usage(key)
                usage = self.key_usage[key]
                
                # Skip exhausted keys
                if usage['is_exhausted']:
                    continue
                
                # Check RPM limit
                current_rpm = len(usage['requests_per_minute'])
                if current_rpm >= self.RPM_LIMIT:
                    logger.debug(f"Key ...{key[-4:]} at RPM limit ({current_rpm}/{self.RPM_LIMIT})")
                    continue
                
                # Check RPD limit
                if usage['requests_today'] >= self.RPD_LIMIT:
                    usage['is_exhausted'] = True
                    logger.warning(f"Key ...{key[-4:]} exhausted for today ({usage['requests_today']}/{self.RPD_LIMIT})")
                    continue
                
                # Check TPM limit
                current_tokens = sum(tokens for _, tokens in usage['tokens_per_minute'])
                if current_tokens + estimated_tokens > self.TPM_LIMIT:
                    logger.debug(f"Key ...{key[-4:]} would exceed TPM limit")
                    continue
                
                # Key is available, record usage
                now = datetime.now()
                usage['requests_per_minute'].append(now)
                usage['requests_today'] += 1
                usage['tokens_per_minute'].append((now, estimated_tokens))
                
                logger.debug(f"Using key ...{key[-4:]} (RPM: {current_rpm+1}/{self.RPM_LIMIT}, RPD: {usage['requests_today']}/{self.RPD_LIMIT})")
                return key
        
        return None
    
    def wait_for_available_key(self, estimated_tokens: int, max_wait: int = 65) -> Optional[str]:
        """Wait for an available key with exponential backoff"""
        start_time = time.time()
        wait_time = 1
        
        while time.time() - start_time < max_wait:
            key = self.get_available_key(estimated_tokens)
            if key:
                return key
            
            logger.info(f"All keys at capacity. Waiting {wait_time}s...")
            time.sleep(wait_time)
            wait_time = min(wait_time * 2, 10)  # Exponential backoff up to 10s
        
        logger.error("No API keys available after waiting")
        return None


# Helper functions for text processing
def deduplicate_lines(lines: List[str], similarity_threshold: float = 0.85) -> List[str]:
    """Remove duplicate lines using fuzzy matching"""
    if not lines:
        return []
    
    unique_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check similarity with existing unique lines
        is_duplicate = False
        for unique_line in unique_lines:
            similarity = SequenceMatcher(None, line.lower(), unique_line.lower()).ratio()
            if similarity >= similarity_threshold:
                is_duplicate = True
                break
        
        if not is_duplicate:
            unique_lines.append(line)
    
    return unique_lines


def fix_ocr_artifacts(text: str) -> str:
    """Fix common OCR errors and artifacts"""
    if not text:
        return text
    
    # Fix common email domain errors
    email_fixes = [
        (r'@gmail\.c(?:om)?(?!\w)', '@gmail.com'),
        (r'@yahoo\.c(?:om)?(?!\w)', '@yahoo.com'),
        (r'@hotmail\.c(?:om)?(?!\w)', '@hotmail.com'),
        (r'@outlook\.c(?:om)?(?!\w)', '@outlook.com'),
        (r'@icloud\.c(?:om)?(?!\w)', '@icloud.com'),
    ]
    
    for pattern, replacement in email_fixes:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # Fix common character substitutions
    text = text.replace('|', 'I')  # Pipe often misread as I
    text = re.sub(r'(?<=[a-zA-Z])0(?=[a-zA-Z])', 'O', text)  # 0 as O in words
    text = re.sub(r'(?<=[a-zA-Z])1(?=[a-zA-Z])', 'l', text)  # 1 as l in words
    
    return text


def validate_email(email: str) -> Optional[str]:
    """Validate and clean email addresses"""
    if not email:
        return None
    
    # Clean common OCR errors first
    email = fix_ocr_artifacts(email)
    
    # Email validation regex
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    
    if re.match(email_pattern, email):
        return email.lower()
    
    return None


def validate_phone(phone: str) -> Optional[str]:
    """Validate and clean phone numbers"""
    if not phone:
        return None
    
    # Extract digits (keep + for international)
    digits = re.sub(r'[^\d+]', '', phone)
    
    # Valid phone numbers have 10-15 digits
    if 10 <= len(digits.replace('+', '')) <= 15:
        return phone
    
    return None


class TextExtractor:
    """Handles text extraction from various file formats"""
    
    _ocr_available = None
    
    @classmethod
    def check_ocr_availability(cls):
        """Check if OCR is available (cached)"""
        if cls._ocr_available is None:
            try:
                pytesseract.get_tesseract_version()
                cls._ocr_available = True
            except:
                cls._ocr_available = False
        return cls._ocr_available
    
    # ✅ 2. IMPROVED PDF TEXT EXTRACTION
    @staticmethod
    def extract_text_from_pdf_bytes(file_bytes) -> str:
        """Extract text from PDF with OCR fallback and fuzzy deduplication"""
        text = ""
        
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                logger.info(f"Processing PDF with {len(pdf.pages)} pages")
                
                for page_num, page in enumerate(pdf.pages):
                    # Step 1: Extract text using pdfplumber
                    logger.debug(f"Extracting text from page {page_num + 1}")
                    page_text = page.extract_text() or ""
                    
                    # Step 2: Check if we need OCR fallback (less than 30 chars)
                    if len(page_text.strip()) < 30 and TextExtractor.check_ocr_availability():
                        logger.info(f"Page {page_num + 1} has minimal text ({len(page_text.strip())} chars), applying OCR fallback")
                        
                        try:
                            # Convert page to high-resolution image for OCR
                            page_image = page.to_image(resolution=300).original
                            ocr_text = pytesseract.image_to_string(page_image, lang='eng')
                            
                            # Step 3: Merge texts intelligently
                            if page_text.strip() and ocr_text.strip():
                                # Both have content - merge and deduplicate
                                logger.debug(f"Merging pdfplumber + OCR text for page {page_num + 1}")
                                combined_lines = page_text.split('\n') + ocr_text.split('\n')
                                deduplicated = deduplicate_lines(combined_lines)
                                page_text = '\n'.join(deduplicated)
                                logger.info(f"Page {page_num + 1}: Merged {len(combined_lines)} lines → {len(deduplicated)} unique lines")
                            elif ocr_text.strip():
                                # Only OCR has content
                                page_text = ocr_text
                                logger.debug(f"Using OCR text only for page {page_num + 1}")
                                
                        except Exception as e:
                            logger.warning(f"OCR fallback failed for page {page_num + 1}: {e}")
                    
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
        
        return text
    
    @staticmethod
    def extract_text_from_docx_bytes(file_bytes) -> str:
        """Extract text from DOCX bytes in memory"""
        text = ""
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            if doc.tables:
                text += "\n[TABLES]\n"
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    if table_data:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else None)
                        text += df.to_string(index=False) + "\n"
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
        
        return text
    
    # ✅ 3. IMPROVED IMAGE OCR ACCURACY
    @staticmethod
    def extract_text_from_image_bytes(file_bytes) -> str:
        """Extract text from images with preprocessing for better accuracy"""
        if not TextExtractor.check_ocr_availability():
            logger.warning("OCR not available")
            return ""
        
        try:
            # Load image
            image = Image.open(io.BytesIO(file_bytes))
            logger.info(f"Processing image: {image.size} {image.mode}")
            
            # Convert to grayscale for better OCR
            image = image.convert("L")
            logger.debug("Converted image to grayscale")
            
            # Enlarge image 2x for better OCR on small text
            width, height = image.size
            new_size = (width * 2, height * 2)
            image = image.resize(new_size, Image.Resampling.LANCZOS)
            logger.debug(f"Enlarged image from {(width, height)} to {new_size}")
            
            # Optional: Apply adaptive thresholding for better contrast
            try:
                img_array = np.array(image)
                threshold = np.mean(img_array)
                img_array = np.where(img_array > threshold, 255, 0).astype(np.uint8)
                image = Image.fromarray(img_array)
                logger.debug("Applied adaptive thresholding")
            except Exception as e:
                logger.warning(f"Thresholding failed: {e}")
            
            # Apply OCR with optimized config
            text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')
            logger.info(f"OCR extracted {len(text)} characters")
            
            # Clean OCR artifacts
            text = fix_ocr_artifacts(text)
            
            return text
            
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt_bytes(file_bytes) -> str:
        """Extract text from TXT bytes"""
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"TXT read error: {e}")
            return ""


class SimpleResumeParser:
    """Simple resume parser using Gemini API"""
    
    def __init__(self):
        load_dotenv()
        
        # Initialize key manager
        self.key_manager = GeminiKeyManager()
        
        # Check if we have any keys
        if not self.key_manager.keys:
            logger.warning("No API keys found. Parser will extract text only.")
            self.api_key = None
        else:
            # For backward compatibility
            self.api_key = self.key_manager.keys[0] if self.key_manager.keys else None
        
        # ✅ 4. IMPROVED GEMINI PROMPT
        self.parsing_prompt =  """ You are a specialized resume parsing AI. Your task is to extract structured data from resume text.

CRITICAL RULES:
1. Return ONLY valid JSON - no markdown, no code blocks, no explanations
2. Use null for missing fields, empty arrays [] for missing lists
3. Extract all available information accurately
4. For total_experience_years, CALCULATE from work_experience job durations, NOT profile summaries

EXACT JSON FORMAT REQUIRED:
{
  "name": "full name",
  "email": "email address", 
  "phone": "phone number",
  "linkedin": "LinkedIn URL",
  "github": "GitHub URL",
  "skills": ["skill1", "skill2", ...],
  "ug_education": {
    "degree": "Bachelor's degree name",
    "college": "college/university name", 
    "year": graduation year as number
  },
  "pg_education": {
    "degree": "Master's/PhD degree name",
    "college": "college/university name",
    "year": graduation year as number  
  },
  "total_experience_years": "CALCULATE by summing ALL job durations from work_experience: For each job (end_year - start_year). If end_year is null (current), use 2024. Round to 1 decimal. Return as string like '4.5', '6.0'. If no work history, return null.",
  "work_experience": [
    {
      "title": "job title",
      "company": "company name", 
      "start_year": start year as number,
      "end_year": end year as number or null if current
    }
  ]
}

Resume content between delimiters:
<<<resume>>> {text} <<<end>>>

REMINDER: Output ONLY the JSON object. Nothing else."""
    
    def is_diploma_entry(self, degree, college):
        """
        Check if the education entry is a diploma-level qualification.
        Returns True if it's a diploma (should be filtered out from UG/PG).
        """
        if not degree and not college:
            return False
        
        degree_str = str(degree).lower() if degree else ""
        college_str = str(college).lower() if college else ""
        
        return "diploma" in degree_str or "diploma" in college_str or "polytechnic" in college_str
    
    def _load_api_key(self) -> Optional[str]:
        """Load API key from environment (kept for compatibility)"""
        # This method is kept for backward compatibility
        # Actual key selection is now handled by GeminiKeyManager
        if self.key_manager and self.key_manager.keys:
            return self.key_manager.keys[0]
        
        # Fallback to original logic
        key = os.getenv('GEMINI_API_KEY_1')
        if key:
            return key
        
        key = os.getenv('GEMINI_API_KEY')
        return key
    
    def extract_text(self, file_obj, filename: str) -> str:
        """Extract text from file bytes with deduplication"""
        ext = Path(filename).suffix.lower()
        
        if ext == '.pdf':
            text = TextExtractor.extract_text_from_pdf_bytes(file_obj)
        elif ext == '.docx':
            text = TextExtractor.extract_text_from_docx_bytes(file_obj)
        elif ext == '.txt':
            text = TextExtractor.extract_text_from_txt_bytes(file_obj)
        elif ext in {'.png', '.jpg', '.jpeg', '.tiff', '.bmp'}:
            text = TextExtractor.extract_text_from_image_bytes(file_obj)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Deduplicate lines to reduce token count
        lines = text.split('\n')
        deduplicated = deduplicate_lines(lines)
        return '\n'.join(deduplicated)
    
    def parse_with_gemini(self, text_content: str) -> Optional[Dict]:
        """Parse resume text with Gemini API using key rotation"""
        if not self.key_manager.keys:
            return None
        
        try:
            # Prepare prompt
            prompt = self.parsing_prompt.replace("{text}", text_content)
            
            # Estimate tokens
            estimated_tokens = self.key_manager.estimate_tokens(prompt)
            logger.debug(f"Estimated tokens for request: {estimated_tokens}")
            
            # Get available key with rate limit checking
            api_key = self.key_manager.wait_for_available_key(estimated_tokens)
            
            if not api_key:
                logger.error("No API keys available due to rate limits")
                return None
            
            # Configure and call Gemini
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(prompt)
            
            return self._parse_json_response(response.text)
            
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return None
    
    # ✅ 5. ENHANCED JSON PARSING WITH VALIDATION AND EXPERIENCE NORMALIZATION
    def _parse_json_response(self, json_text: str) -> Dict[str, Any]:
        """Parse and validate JSON response from Gemini"""
        try:
            # Clean JSON response
            if "```json" in json_text:
                json_text = json_text.split("```json")[1].split("```")[0]
            elif "```" in json_text:
                json_text = json_text.split("```")[1].split("```")[0]
            
            parsed = json.loads(json_text.strip())
            
            # Helper function to normalize experience values
            def normalize_experience(exp_value):
                """Normalize experience value to a consistent string format"""
                if not exp_value:
                    return None
                
                exp_str = str(exp_value).strip().lower()
                
                # Handle empty or None values
                if not exp_str or exp_str in ['null', 'none', '']:
                    return None
                
                # Keep "X+" format as-is (e.g., "4+" -> "4+")
                if '+' in exp_str:
                    # Extract the number before +
                    match = re.search(r'(\d+(?:\.\d+)?)\s*\+', exp_str)
                    if match:
                        return f"{match.group(1)}+"
                    return None
                
                # Handle "X years Y months" format
                years_months_match = re.search(r'(\d+(?:\.\d+)?)\s*years?\s*(\d+)\s*months?', exp_str)
                if years_months_match:
                    years = float(years_months_match.group(1))
                    months = float(years_months_match.group(2))
                    # Convert months to decimal (3 months = 0.2, 6 months = 0.5, etc.)
                    total_years = years + round(months / 12, 1)
                    return str(total_years)
                
                # Handle "X months" only format
                months_only_match = re.search(r'(\d+(?:\.\d+)?)\s*months?', exp_str)
                if months_only_match:
                    months = float(months_only_match.group(1))
                    years = round(months / 12, 1)
                    return str(years)
                
                # Handle "X years" format (extract just the number)
                years_only_match = re.search(r'(\d+(?:\.\d+)?)\s*years?', exp_str)
                if years_only_match:
                    return str(float(years_only_match.group(1)))
                
                # Handle plain numbers (e.g., "6", "4.5")
                number_match = re.search(r'^(\d+(?:\.\d+)?)$', exp_str)
                if number_match:
                    value = float(number_match.group(1))
                     # Return integer format if it's a whole number
                    if value == int(value):
                        return str(int(value))
                    else:
                        return str(value)
                
                # If we can't parse it, return None
                return None
            
            # Build result with validation and cleanup
            result = {
                'name': parsed.get('name'),
                'email': validate_email(parsed.get('email')),
                'phone': validate_phone(parsed.get('phone')),
                'linkedin': parsed.get('linkedin'),
                'github': parsed.get('github'),
                'skills': parsed.get('skills', []),
                'total_experience_years': normalize_experience(parsed.get('total_experience_years'))
            }
            
            # Handle UG education with diploma filtering
            ug_edu = parsed.get('ug_education', {})
            if ug_edu and isinstance(ug_edu, dict):
                ug_degree = ug_edu.get('degree')
                ug_college = ug_edu.get('college')
                ug_year = ug_edu.get('year')
                
                # Filter out diploma entries
                if self.is_diploma_entry(ug_degree, ug_college):
                    logger.info(f"Filtering out diploma from UG education: {ug_degree} at {ug_college}")
                    result['ug_degree'] = None
                    result['ug_college'] = None
                    result['ug_year'] = None
                else:
                    result['ug_degree'] = ug_degree
                    result['ug_college'] = ug_college
                    result['ug_year'] = ug_year
            else:
                result['ug_degree'] = None
                result['ug_college'] = None
                result['ug_year'] = None
            
            # Handle PG education with diploma filtering
            pg_edu = parsed.get('pg_education', {})
            if pg_edu and isinstance(pg_edu, dict):
                pg_degree = pg_edu.get('degree')
                pg_college = pg_edu.get('college')
                pg_year = pg_edu.get('year')
                
                # Filter out diploma entries
                if self.is_diploma_entry(pg_degree, pg_college):
                    logger.info(f"Filtering out diploma from PG education: {pg_degree} at {pg_college}")
                    result['pg_degree'] = None
                    result['pg_college'] = None
                    result['pg_year'] = None
                else:
                    result['pg_degree'] = pg_degree
                    result['pg_college'] = pg_college
                    result['pg_year'] = pg_year
            else:
                result['pg_degree'] = None
                result['pg_college'] = None
                result['pg_year'] = None
            
            # Handle work experience
            result['work_experience'] = parsed.get('work_experience', [])
            
            return result
                
        except Exception as e:
            logger.error(f"JSON parsing error: {e}")
            return None
    
    def process_resume(self, file_obj, filename: str) -> Dict:
        """Process a single resume"""
        result = {
            'filename': filename,
            'success': False,
            'data': None,
            'error': None
        }
        
        try:
            # Extract text
            text_content = self.extract_text(file_obj, filename)
            
            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("No meaningful text extracted")
            
            # Parse with Gemini if available
            if self.key_manager.keys:
                parsed_data = self.parse_with_gemini(text_content)
                
                if parsed_data:
                    result['data'] = parsed_data
                    result['success'] = True
                else:
                    result['error'] = "Failed to parse resume"
            else:
                # If no API key, just return success with empty data
                result['data'] = {
                    'name': None,
                    'email': None,
                    'phone': None,
                    'linkedin': None,
                    'github': None,
                    'skills': [],
                    'ug_degree': None,
                    'ug_college': None,
                    'ug_year': None,
                    'pg_degree': None,
                    'pg_college': None,
                    'pg_year': None,
                    'total_experience_years': None,
                    'work_experience': []
                }
                result['success'] = True
                result['error'] = "No API key - text extracted but not parsed"
                
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"Error processing {filename}: {e}")
        
        return result